"""
Service de génération de documents HERMES par IA (Claude API).
Le CP saisit les informations de base du projet → l'IA rédige les sections
des templates Word selon la méthodologie HERMES 5.1.
"""
import anthropic
import json
from pathlib import Path
from docxtpl import DocxTemplate
from app.core.config import settings

client = anthropic.AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)

# Mapping template_key → fichier template Word
TEMPLATE_MAP = {
    "mandat_initialisation": "templates/mandat_initialisation.docx",
    "mandat_projet": "templates/mandat_projet.docx",
    "planning_jalons": "templates/planning_jalons.docx",
}

# Sections à générer par template
TEMPLATE_SECTIONS = {
    "mandat_initialisation": [
        "situation_initiale",
        "objectifs_initialisation",
        "perimetre",
        "organisation_projet",
        "risques_initiaux",
        "criteres_succes",
    ],
    "mandat_projet": [
        "situation_initiale",
        "objectifs_projet",
        "perimetre_exclusions",
        "organisation_projet",
        "planning_jalons",
        "risques_identifies",
        "budget_ressources_note",
        "criteres_succes",
    ],
    "planning_jalons": [
        "jalons_liste",
        "dependances",
        "ressources_critiques",
    ],
}

HERMES_SYSTEM_PROMPT = """Tu es un expert en gestion de projet selon la méthodologie HERMES 5.1 (standard suisse).
Tu rédiges des livrables officiels HERMES en français, dans un style professionnel, clair et concis.
Tes réponses doivent être directement utilisables dans un document Word officiel.
Respecte strictement la terminologie HERMES (commanditaire, chef de projet, jalons, livrables, etc.).
Chaque section doit être autonome et complète — 2 à 5 phrases, factuelle et structurée."""


async def generate_hermes_document(deliverable_id: int, template_key: str) -> dict:
    """
    Génère un document HERMES en remplissant les sections via Claude.
    Retourne un dict avec les sections générées et le chemin du fichier Word.
    """
    # 1. Récupérer les données du projet depuis la DB
    project_data = await _get_project_data(deliverable_id)

    # 2. Générer chaque section avec Claude
    sections = {}
    for section_key in TEMPLATE_SECTIONS.get(template_key, []):
        sections[section_key] = await _generate_section(
            section_key, template_key, project_data
        )

    # 3. Fusionner dans le template Word
    output_path = await _merge_into_template(template_key, project_data, sections, deliverable_id)

    # 4. Mettre à jour la DB
    await _update_deliverable(deliverable_id, output_path, len(sections))

    return {"sections": sections, "output_path": str(output_path)}


async def _generate_section(section_key: str, template_key: str, project_data: dict) -> str:
    """Appelle Claude pour générer une section spécifique."""
    prompt = _build_section_prompt(section_key, template_key, project_data)

    message = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=800,
        system=HERMES_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text.strip()


def _build_section_prompt(section_key: str, template_key: str, project_data: dict) -> str:
    section_labels = {
        "situation_initiale": "Situation initiale (contexte, problèmes à résoudre)",
        "objectifs_projet": "Objectifs du projet (SMART, 3-5 objectifs numérotés)",
        "objectifs_initialisation": "Objectifs de la phase d'initialisation",
        "perimetre": "Périmètre du projet (ce qui est inclus et exclu)",
        "perimetre_exclusions": "Périmètre et exclusions explicites",
        "organisation_projet": "Organisation du projet (rôles et responsabilités HERMES)",
        "planning_jalons": "Planning et jalons principaux (tableau avec dates)",
        "risques_initiaux": "Risques initiaux identifiés (liste avec probabilité/impact/mitigation)",
        "risques_identifies": "Risques identifiés et mesures de mitigation",
        "criteres_succes": "Critères de succès mesurables du projet",
        "budget_ressources_note": "Note sur le budget et les ressources (sans chiffres précis si non fournis)",
        "jalons_liste": "Liste des jalons avec dates et responsables",
        "dependances": "Dépendances entre jalons et livrables",
        "ressources_critiques": "Ressources critiques et points d'attention",
    }

    label = section_labels.get(section_key, section_key)
    return f"""Rédige la section "{label}" pour le document HERMES "{template_key.replace('_', ' ').title()}".

Données du projet :
- Nom : {project_data.get('name', 'N/A')}
- Client / Commanditaire : {project_data.get('client', 'N/A')}
- Chef de projet : {project_data.get('project_manager', 'N/A')}
- Budget estimé : CHF {project_data.get('budget_chf', 'N/A')}
- Délai souhaité : {project_data.get('end_date_planned', 'N/A')}
- Description : {project_data.get('description', 'N/A')}

Rédige uniquement le contenu de cette section, sans titre ni balise Markdown."""


async def _get_project_data(deliverable_id: int) -> dict:
    """Récupère les données du projet lié au livrable depuis la DB."""
    from app.core.database import AsyncSessionLocal
    from app.models.project import Deliverable, Project
    from sqlmodel import select

    async with AsyncSessionLocal() as session:
        result = await session.execute(
            select(Deliverable).where(Deliverable.id == deliverable_id)
        )
        d = result.scalar_one_or_none()
        if not d:
            return {}
        project = await session.get(Project, d.project_id)
        if not project:
            return {}
        return project.model_dump()


async def _merge_into_template(
    template_key: str, project_data: dict, sections: dict, deliverable_id: int
) -> Path:
    """Fusionne les sections générées dans le template Word."""
    template_path = Path(__file__).parent.parent / TEMPLATE_MAP.get(
        template_key, "templates/mandat_projet.docx"
    )

    output_dir = Path("/app/uploads/generated")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{template_key}_{deliverable_id}.docx"

    if template_path.exists():
        doc = DocxTemplate(template_path)
        context = {**project_data, **sections}
        doc.render(context)
        doc.save(output_path)
    else:
        # Mode fallback sans template : créer un docx simple
        from docx import Document
        doc = Document()
        doc.add_heading(f"HERMES — {template_key.replace('_', ' ').title()}", 0)
        doc.add_paragraph(f"Projet : {project_data.get('name', '')}")
        for key, content in sections.items():
            doc.add_heading(key.replace("_", " ").title(), level=1)
            doc.add_paragraph(content)
        doc.save(output_path)

    return output_path


async def _update_deliverable(deliverable_id: int, output_path: Path, sections_done: int):
    from app.core.database import AsyncSessionLocal
    from app.models.project import Deliverable
    from datetime import datetime

    async with AsyncSessionLocal() as session:
        d = await session.get(Deliverable, deliverable_id)
        if d:
            d.ai_generated = True
            d.ai_sections_complete = sections_done
            d.document_path = str(output_path)
            d.updated_at = datetime.utcnow()
            session.add(d)
            await session.commit()
