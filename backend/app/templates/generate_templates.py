"""
Script à exécuter UNE SEULE FOIS pour générer les templates Word HERMES avec les placeholders docxtpl.
Usage : python generate_templates.py  (depuis le dossier backend/app/templates/)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


BLUE = RGBColor(0x1E, 0x40, 0xAF)   # blue-800
GRAY = RGBColor(0x6B, 0x72, 0x80)   # gray-500


def add_cover(doc: Document, title: str, subtitle: str):
    doc.add_heading("HERMES 5.1", level=1).runs[0].font.color.rgb = BLUE
    t = doc.add_heading(title, level=2)
    t.runs[0].font.color.rgb = BLUE

    p = doc.add_paragraph(subtitle)
    p.runs[0].font.color.rgb = GRAY

    # Infos projet auto-remplies
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Projet :", "{{name}}"),
        ("Client / Commanditaire :", "{{client}}"),
        ("Chef de projet :", "{{project_manager}}"),
        ("Date de création :", "{{created_at}}"),
        ("Statut :", "{{status}}"),
    ]
    for i, (label, val) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
    doc.add_page_break()


def add_section(doc: Document, title: str, placeholder: str, hint: str = ""):
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = BLUE
    if hint:
        p = doc.add_paragraph(hint)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = GRAY
    # Placeholder docxtpl — sera remplacé par le contenu IA
    body = doc.add_paragraph(f"{{{{{placeholder}}}}}")
    body.paragraph_format.space_after = Pt(12)


# ─── Mandat d'initialisation ───────────────────────────────────────
def create_mandat_initialisation():
    doc = Document()
    add_cover(doc, "Mandat d'initialisation", "Document HERMES — Phase Initialisation")

    add_section(doc, "1. Situation initiale", "situation_initiale",
                "Contexte organisationnel, problèmes à résoudre, opportunités identifiées.")
    add_section(doc, "2. Objectifs de la phase d'initialisation", "objectifs_initialisation",
                "Ce que la phase d'initialisation doit produire comme résultats.")
    add_section(doc, "3. Périmètre", "perimetre",
                "Ce qui est inclus dans le projet et ce qui en est explicitement exclu.")
    add_section(doc, "4. Organisation du projet", "organisation_projet",
                "Rôles HERMES : commanditaire, chef de projet, équipe, instances de décision.")
    add_section(doc, "5. Risques initiaux", "risques_initiaux",
                "Risques identifiés avec probabilité, impact et mesures de mitigation.")
    add_section(doc, "6. Critères de succès", "criteres_succes",
                "Indicateurs mesurables permettant de juger la réussite du projet.")

    out = Path(__file__).parent / "mandat_initialisation.docx"
    doc.save(out)
    print(f"✅  {out}")


# ─── Mandat de projet ─────────────────────────────────────────────
def create_mandat_projet():
    doc = Document()
    add_cover(doc, "Mandat de projet", "Document HERMES — Phase Conception")

    add_section(doc, "1. Situation initiale", "situation_initiale",
                "Contexte stratégique justifiant le lancement du projet.")
    add_section(doc, "2. Objectifs du projet", "objectifs_projet",
                "Objectifs SMART numérotés (3 à 5 objectifs mesurables).")
    add_section(doc, "3. Périmètre et exclusions", "perimetre_exclusions",
                "Définition précise du périmètre et liste des exclusions explicites.")
    add_section(doc, "4. Organisation du projet", "organisation_projet",
                "Structure organisationnelle, rôles et responsabilités HERMES.")
    add_section(doc, "5. Planning et jalons", "planning_jalons",
                "Calendrier prévisionnel avec jalons principaux et dates cibles.")
    add_section(doc, "6. Risques identifiés", "risques_identifies",
                "Analyse des risques avec probabilité, impact et plan de mitigation.")
    add_section(doc, "7. Budget et ressources", "budget_ressources_note",
                "Vue d'ensemble budgétaire et ressources nécessaires.")
    add_section(doc, "8. Critères de succès", "criteres_succes",
                "Indicateurs de performance et critères d'acceptation du projet.")

    out = Path(__file__).parent / "mandat_projet.docx"
    doc.save(out)
    print(f"✅  {out}")


# ─── Planning avec jalons ─────────────────────────────────────────
def create_planning_jalons():
    doc = Document()
    add_cover(doc, "Planning avec jalons", "Document HERMES — Vue planning")

    add_section(doc, "1. Liste des jalons", "jalons_liste",
                "Jalons HERMES avec dates, responsables et livrables associés.")
    add_section(doc, "2. Dépendances", "dependances",
                "Liens de dépendance entre jalons et livrables critiques.")
    add_section(doc, "3. Ressources critiques", "ressources_critiques",
                "Ressources et compétences critiques nécessaires au respect du planning.")

    # Tableau d'exemple
    doc.add_heading("4. Tableau récapitulatif", level=1).runs[0].font.color.rgb = BLUE
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Jalon", "Description", "Date prévue", "Responsable", "Statut"]):
        hdr[i].text = h

    out = Path(__file__).parent / "planning_jalons.docx"
    doc.save(out)
    print(f"✅  {out}")


if __name__ == "__main__":
    create_mandat_initialisation()
    create_mandat_projet()
    create_planning_jalons()
    print("\nTemplates HERMES générés avec succès.")
